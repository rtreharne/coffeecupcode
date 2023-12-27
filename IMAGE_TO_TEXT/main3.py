import pytesseract
from PIL import Image
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load the image
image_path = "page_1.jpg"
image = Image.open(image_path)

# Extract text using pytesseract
text = pytesseract.image_to_string(image)

# Create a new Word document
doc = Document()

# Add paragraphs with preserved formatting
for line in text.splitlines():
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(line)

    # Preserve italics formatting
    if line.startswith("_") and line.endswith("_"):
        run.italic = True
        line = line.strip("_")

    # Set font size and alignment
    run.font.size = Pt(12)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Save the document
doc.save("output2.docx")
