import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

# URL of the webpage to extract text from
url = "https://www.gutenberg.org/files/21374/21374-h/21374-h.htm"

# Fetch the HTML content from the URL
response = requests.get(url)
html_content = response.text

# Parse the HTML content and extract the text
soup = BeautifulSoup(html_content, "html.parser")
text = soup.get_text()

# Split the text into paragraphs
paragraphs = text.split('\n')

# Create a new document
doc = Document()

# Set the default font size and spacing
doc.styles['Normal'].font.size = Pt(10)
doc.styles['Normal'].paragraph_format.line_spacing = 1.5

# Add each paragraph as a separate paragraph in the document
for paragraph in paragraphs:
    doc.add_paragraph(paragraph)

    # Add space between paragraphs
    doc.add_paragraph()

# Save the document as a DOCX file
doc.save("gutenberg.docx")

# Create a new document
doc = Document()

# Set the default font size and spacing
doc.styles['Normal'].font.size = Pt(10)
doc.styles['Normal'].paragraph_format.line_spacing = 1.5

# Add the extracted text to the document
doc.add_paragraph(text)

# Save the document as a DOCX file
doc.save("gutenberg.docx")
