import cv2
import pytesseract
from docx import Document
from docx.shared import Pt

# Load the image
image = cv2.imread('page_1.jpg')

# Check if the image was loaded correctly
if image is None:
    print('Could not open or find the image')
    exit
else:
    # Convert the image to grayscale
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

gray = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]

# Perform OCR using pytesseract
text = pytesseract.image_to_string(gray)

# Get the italicized text and font size
italicized_text = []
font_sizes = []

# Iterate through each character in the text
osd_data = pytesseract.image_to_osd(gray, output_type=pytesseract.Output.DICT)
if 'text' in osd_data:
    for char_info in osd_data['text']:
        if char_info['italic'] == '1':
            italicized_text.append(char_info['char'])
            font_sizes.append(char_info['size'])

# Create a new Word document
doc = Document()

# Add the extracted text to the document
doc.add_heading('Extracted Text', level=1)
doc.add_paragraph(text)

# Add the italicized text to the document
doc.add_heading('Italicized Text', level=1)
doc.add_paragraph(''.join(italicized_text))

# Add the font sizes to the document
doc.add_heading('Font Sizes', level=1)
for size in font_sizes:
    p = doc.add_paragraph()
    run = p.add_run(str(size))
    font = run.font
    font.size = Pt(12)  # Set the font size to 12pt

# Save the document
doc.save('output.docx')
